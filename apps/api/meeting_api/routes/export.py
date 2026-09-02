from __future__ import annotations

import re
from io import BytesIO
from typing import Literal

from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fastapi import APIRouter, HTTPException, Request, Response, status
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.orm import Session

from meeting_api.models import (
    CleanedTranscriptBlock,
    Meeting,
    TranscriptSegment,
)
from meeting_api.speaker_labels import public_speaker_labels
from meeting_api.storage import meeting_dir
from meeting_api.transcript_cleaning import apply_cleaned_blocks
from meeting_api.transcript_format import (
    TranscriptBlock,
    build_transcript_blocks,
    render_transcript_blocks,
)
from meeting_domain import MeetingState

router = APIRouter(prefix="/api/meetings")

MARKDOWN_MEDIA_TYPE = "text/markdown"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TRANSCRIPT_EXPORT_STATES = {
    MeetingState.AWAITING_SPEAKER_REVIEW.value,
    MeetingState.APPLYING_DECISIONS.value,
    MeetingState.GENERATING_MINUTES.value,
    MeetingState.READY.value,
    MeetingState.PARTIAL_READY.value,
}
MINUTES_EXPORT_STATES = {
    MeetingState.READY.value,
    MeetingState.PARTIAL_READY.value,
}
INLINE_BOLD_PATTERN = re.compile(r"(\*\*[^*]+\*\*)")
LIST_ITEM_PATTERN = re.compile(r"^(?P<nested>  )?- (?P<content>.*)$")
TASK_PATTERN = re.compile(r"^\[(?P<checked>[ xX])\] (?P<content>.*)$")


class TranscriptBlockResponse(BaseModel):
    start_seconds: float
    end_seconds: float
    # 公开说话人标签：确认后的名字或「说话人 N」。
    label: str
    text: str
    # 该块的清洗文本；清洗失败或哈希对不上时为 None，前端回退原文。
    cleaned_text: str | None


class TranscriptResponse(BaseModel):
    # 块级结构才是接口契约；markdown 字段只给导出预览与旧调用方。
    blocks: list[TranscriptBlockResponse]
    cleaned_available: bool
    raw_markdown: str
    cleaned_markdown: str | None


def _get_meeting(session: Session, meeting_id: str) -> Meeting:
    meeting = session.get(Meeting, meeting_id)
    if meeting is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="会议不存在")
    return meeting


def _attachment_headers(meeting_id: str, suffix: str) -> dict[str, str]:
    return {"Content-Disposition": f'attachment; filename="meeting-{meeting_id}-{suffix}"'}


def _build_export_transcript(
    session: Session,
    meeting_id: str,
    variant: Literal["raw", "cleaned"] = "raw",
) -> str:
    blocks = _build_transcript_blocks(session, meeting_id)
    if variant == "cleaned":
        # 哈希对不上的块自动落回原文，所以 cleaned 请求永远给得出内容。
        blocks, _ = apply_cleaned_blocks(
            blocks,
            _cleaned_rows_by_index(session, meeting_id),
        )
    return _render_transcript_markdown(blocks)


def _build_cleaned_export_transcript(
    session: Session,
    meeting_id: str,
) -> str | None:
    blocks = _build_transcript_blocks(session, meeting_id)
    cleaned_blocks, changed = apply_cleaned_blocks(
        blocks,
        _cleaned_rows_by_index(session, meeting_id),
    )
    if not changed:
        return None
    return _render_transcript_markdown(cleaned_blocks)


def _build_transcript_blocks(session: Session, meeting_id: str) -> list[TranscriptBlock]:
    segments = session.scalars(
        select(TranscriptSegment)
        .where(TranscriptSegment.meeting_id == meeting_id)
        .order_by(TranscriptSegment.start_seconds, TranscriptSegment.id)
    ).all()
    labels = public_speaker_labels(session, meeting_id)
    return build_transcript_blocks(
        [
            (
                segment.start_seconds,
                segment.end_seconds,
                labels.get(segment.cluster_id, "说话人"),
                segment.text,
            )
            for segment in segments
        ]
    )


def _cleaned_rows_by_index(
    session: Session,
    meeting_id: str,
) -> dict[int, tuple[str, str]]:
    rows = session.scalars(
        select(CleanedTranscriptBlock).where(
            CleanedTranscriptBlock.meeting_id == meeting_id
        )
    ).all()
    return {
        row.block_index: (row.raw_sha256, row.cleaned_text)
        for row in rows
    }


def _render_transcript_markdown(blocks: list[TranscriptBlock]) -> str:
    return "\n".join(["# 会议转写", "", render_transcript_blocks(blocks)])


def _read_minutes(request: Request, meeting_id: str) -> str:
    minutes_path = meeting_dir(request.app.state.settings, meeting_id) / "minutes.md"
    if not minutes_path.is_file():
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="纪要尚未生成",
        )
    return minutes_path.read_text(encoding="utf-8")


def _require_export_state(meeting: Meeting, allowed_states: set[str]) -> None:
    if meeting.state not in allowed_states:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="会议当前状态不可导出该文件",
        )


def _add_markdown_runs(paragraph: Paragraph, text: str) -> None:
    for part in INLINE_BOLD_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def _render_minutes_docx(document: DocumentObject, markdown: str) -> None:
    previous: Paragraph | None = None
    for line in markdown.splitlines():
        if not line.strip():
            if previous is not None:
                previous.paragraph_format.space_after = Pt(8)
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 2")
            content = line[3:]
        elif line.startswith("# "):
            paragraph = document.add_paragraph(style="Heading 1")
            content = line[2:]
        elif match := LIST_ITEM_PATTERN.match(line):
            paragraph = document.add_paragraph(
                style="List Bullet 2" if match.group("nested") else "List Bullet"
            )
            content = match.group("content")
            if task := TASK_PATTERN.match(content):
                checkbox = "☑" if task.group("checked").lower() == "x" else "☐"
                paragraph.add_run(f"{checkbox} ")
                content = task.group("content")
        else:
            paragraph = document.add_paragraph()
            content = line

        _add_markdown_runs(paragraph, content)
        previous = paragraph


@router.get("/{meeting_id}/export/transcript.md")
def export_transcript(
    meeting_id: str,
    request: Request,
    variant: Literal["raw", "cleaned"] = "raw",
) -> Response:
    with request.app.state.session_factory() as session:
        meeting = _get_meeting(session, meeting_id)
        _require_export_state(meeting, TRANSCRIPT_EXPORT_STATES)
        markdown = _build_export_transcript(session, meeting_id, variant=variant)
    return Response(
        content=markdown,
        media_type=MARKDOWN_MEDIA_TYPE,
        headers=_attachment_headers(meeting_id, "transcript.md"),
    )


@router.get("/{meeting_id}/transcript", response_model=TranscriptResponse)
def get_transcript(meeting_id: str, request: Request) -> TranscriptResponse:
    with request.app.state.session_factory() as session:
        meeting = _get_meeting(session, meeting_id)
        _require_export_state(meeting, TRANSCRIPT_EXPORT_STATES)
        blocks = _build_transcript_blocks(session, meeting_id)
        cleaned_blocks, changed = apply_cleaned_blocks(
            blocks, _cleaned_rows_by_index(session, meeting_id)
        )
        return TranscriptResponse(
            blocks=[
                TranscriptBlockResponse(
                    start_seconds=block.start_seconds,
                    end_seconds=block.end_seconds,
                    label=block.label,
                    text=block.text,
                    cleaned_text=cleaned.text if cleaned.text != block.text else None,
                )
                for block, cleaned in zip(blocks, cleaned_blocks, strict=True)
            ],
            cleaned_available=changed,
            raw_markdown=_render_transcript_markdown(blocks),
            cleaned_markdown=_render_transcript_markdown(cleaned_blocks) if changed else None,
        )


@router.get("/{meeting_id}/export/minutes.md")
def export_minutes_markdown(meeting_id: str, request: Request) -> Response:
    with request.app.state.session_factory() as session:
        meeting = _get_meeting(session, meeting_id)
        _require_export_state(meeting, MINUTES_EXPORT_STATES)
    markdown = _read_minutes(request, meeting_id)
    return Response(
        content=markdown,
        media_type=MARKDOWN_MEDIA_TYPE,
        headers=_attachment_headers(meeting_id, "minutes.md"),
    )


@router.get("/{meeting_id}/export/minutes.docx")
def export_minutes_docx(meeting_id: str, request: Request) -> Response:
    with request.app.state.session_factory() as session:
        meeting = _get_meeting(session, meeting_id)
        _require_export_state(meeting, MINUTES_EXPORT_STATES)
    markdown = _read_minutes(request, meeting_id)

    document = Document()
    _render_minutes_docx(document, markdown)
    output = BytesIO()
    document.save(output)
    return Response(
        content=output.getvalue(),
        media_type=DOCX_MEDIA_TYPE,
        headers=_attachment_headers(meeting_id, "minutes.docx"),
    )
